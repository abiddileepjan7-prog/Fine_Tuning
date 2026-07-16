"""
=========================================================
DOCX Processing → Structured JSON + Timing Report
Libraries:
1. python-docx
2. mammoth
=========================================================
"""

import json
import time
from docx import Document
import mammoth

FILE = "file-sample_100kB.docx"
OUTPUT_JSON = "docx_structured_output.json"
TIMING_JSON = "docx_timing_report.json"

# ==========================================================
# Unified output schema
# ==========================================================

result = {
    "source_file": FILE,
    "metadata": {},
    "paragraphs": [],
    "tables": [],
    "html": ""
}

# ==========================================================
# Timing tracker
# ==========================================================

timings = {}


def time_it(label, func):
    """Runs func(), records elapsed time, returns func's result."""

    start = time.perf_counter()
    output = func()
    end = time.perf_counter()

    elapsed = round(end - start, 4)

    timings[label] = {"time_seconds": elapsed}

    print(f"\n[TIMER] {label} took {elapsed} seconds")

    return output


# ==========================================================
# 1. PYTHON-DOCX — paragraphs, metadata, tables
# ==========================================================

print("=" * 70)
print("1. PYTHON-DOCX")
print("=" * 70)


def run_python_docx():

    doc = Document(FILE)

    # ------------------------------------------------------
    # Paragraphs
    # ------------------------------------------------------

    total_chars = 0

    for para in doc.paragraphs:

        text = para.text.strip()

        if text:
            result["paragraphs"].append(text)
            total_chars += len(text)

    # ------------------------------------------------------
    # Metadata
    # ------------------------------------------------------

    meta = doc.core_properties

    result["metadata"] = {
        "author": meta.author,
        "title": meta.title,
        "subject": meta.subject,
        "created": str(meta.created) if meta.created else None,
        "modified": str(meta.modified) if meta.modified else None,
    }

    # ------------------------------------------------------
    # Tables
    # ------------------------------------------------------

    for table in doc.tables:

        rows = []

        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])

        result["tables"].append(rows)

    return total_chars


chars_docx = time_it("python_docx", run_python_docx)

timings["python_docx"]["paragraphs_extracted"] = len(result["paragraphs"])
timings["python_docx"]["chars_extracted"] = chars_docx
timings["python_docx"]["tables_found"] = len(result["tables"])

print("\nParagraphs Extracted :", len(result["paragraphs"]))
print("Metadata :", result["metadata"])
print("Tables Extracted :", len(result["tables"]))

# ==========================================================
# 2. MAMMOTH — DOCX to structured HTML
# ==========================================================

print("\n")
print("=" * 70)
print("2. MAMMOTH")
print("=" * 70)


def run_mammoth():

    with open(FILE, "rb") as f:
        conversion = mammoth.convert_to_html(f)

    result["html"] = conversion.value

    if conversion.messages:
        print("\nMessages")
        for msg in conversion.messages:
            print(msg)

    return len(conversion.value)


html_length = time_it("mammoth", run_mammoth)

timings["mammoth"]["html_length"] = html_length

print("\nHTML Length :", html_length)

# ==========================================================
# SAVE FINAL STRUCTURED JSON
# ==========================================================

with open(OUTPUT_JSON, "w", encoding="utf-8") as f:
    json.dump(result, f, indent=2, ensure_ascii=False)

# ==========================================================
# TIMING REPORT — print to terminal + save to file
# ==========================================================

timing_report = {
    "source_file": FILE,
    "timings": timings
}

print("\n")
print("=" * 70)
print("TIMING REPORT")
print("=" * 70)

print(f"\n{'Library':<18}{'Time (s)'}")
print("-" * 30)

for lib, data in timings.items():
    print(f"{lib:<18}{data['time_seconds']}")

print("\nFull Timing Report (JSON)")
print(json.dumps(timing_report, indent=2, ensure_ascii=False))

with open(TIMING_JSON, "w", encoding="utf-8") as f:
    json.dump(timing_report, f, indent=2, ensure_ascii=False)

print(f"\nTiming report saved to {TIMING_JSON}")

print("\nDone.")